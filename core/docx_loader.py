"""DOCX text and table extraction for InsightFlow AI.

The module separates narrative document content from tabular content:
- Paragraphs and headings become evidence-bearing text blocks.
- Word tables become Pandas DataFrames.
- One table can be selected for the existing analytics pipeline.

No LLM is used during extraction.
"""

from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, BinaryIO

import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class DocumentTextBlock:
    """One evidence-bearing paragraph extracted from a DOCX file."""

    block_id: str
    paragraph_index: int
    text: str
    style: str
    block_type: str
    heading_level: int | None
    section_path: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class DocumentTableData:
    """One extracted Word table and its analytical metadata."""

    table_index: int
    table_id: str
    title: str
    row_count: int
    column_count: int
    headers: list[str]
    dataframe: pd.DataFrame
    source_reference: str
    warnings: list[str] = field(default_factory=list)

    def metadata(self) -> dict[str, Any]:
        data = asdict(self)
        data.pop("dataframe", None)
        return data


@dataclass
class DocxMetadata:
    """Document-level extraction metadata."""

    filename: str
    format: str
    paragraph_count: int
    non_empty_paragraph_count: int
    heading_count: int
    table_count: int
    extracted_table_count: int
    text_character_count: int
    title: str | None
    author: str | None
    subject: str | None
    created: str | None
    modified: str | None
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ExtractedDocx:
    """Complete deterministic extraction result for one DOCX file."""

    metadata: DocxMetadata
    text_blocks: list[DocumentTextBlock]
    tables: list[DocumentTableData]

    def table_summaries(self) -> list[dict[str, Any]]:
        return [table.metadata() for table in self.tables]

    def text_records(self) -> list[dict[str, Any]]:
        return [block.to_dict() for block in self.text_blocks]

    def combined_text(self) -> str:
        return "\n\n".join(
            block.text
            for block in self.text_blocks
            if block.text.strip()
        )


@dataclass
class SelectedDocxTable:
    """A DOCX table selected for the existing analytics pipeline."""

    dataframe: pd.DataFrame
    table_metadata: dict[str, Any]
    document_metadata: dict[str, Any]


SUPPORTED_EXTENSION = ".docx"


def extract_docx(
    uploaded_file: BinaryIO,
    filename: str,
) -> ExtractedDocx:
    """Extract text, headings, metadata, and tables from a DOCX file."""

    cleaned_filename = _validate_filename(filename)
    raw_bytes = _read_bytes(uploaded_file)

    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as error:
        raise ValueError(
            "Unable to open the DOCX file. Confirm that the file is a "
            f"valid Word document. Technical details: {error}"
        ) from error

    text_blocks = _extract_text_blocks(document)
    tables = _extract_tables(document)
    metadata = _extract_metadata(
        document=document,
        filename=cleaned_filename,
        text_blocks=text_blocks,
        tables=tables,
    )

    return ExtractedDocx(
        metadata=metadata,
        text_blocks=text_blocks,
        tables=tables,
    )


def list_docx_tables(
    uploaded_file: BinaryIO,
    filename: str,
) -> list[dict[str, Any]]:
    """Return table summaries without requiring callers to inspect DataFrames."""

    return extract_docx(uploaded_file, filename).table_summaries()


def select_docx_table(
    uploaded_file: BinaryIO,
    filename: str,
    table_index: int = 0,
) -> SelectedDocxTable:
    """Select one extracted DOCX table for Pandas analytics."""

    extracted = extract_docx(uploaded_file, filename)

    if not extracted.tables:
        raise ValueError(
            "The DOCX file contains no usable tables for tabular analysis."
        )

    if table_index < 0 or table_index >= len(extracted.tables):
        raise ValueError(
            f"Table index {table_index} is outside the valid range. "
            f"Available table indexes: 0 to {len(extracted.tables) - 1}."
        )

    selected = extracted.tables[table_index]

    return SelectedDocxTable(
        dataframe=selected.dataframe.copy(),
        table_metadata=selected.metadata(),
        document_metadata=extracted.metadata.to_dict(),
    )


def _validate_filename(filename: str) -> str:
    """Validate the DOCX filename."""

    if not isinstance(filename, str):
        raise TypeError("filename must be a string")

    cleaned = filename.strip()

    if not cleaned:
        raise ValueError("The uploaded filename is empty.")

    extension = Path(cleaned).suffix.casefold()

    if extension != SUPPORTED_EXTENSION:
        raise ValueError(
            f"Unsupported document type '{extension or 'unknown'}'. "
            "This loader supports DOCX files only."
        )

    return cleaned


def _read_bytes(uploaded_file: BinaryIO) -> bytes:
    """Read all bytes while preserving a reusable seekable stream."""

    if uploaded_file is None:
        raise ValueError("No DOCX file was provided.")

    try:
        uploaded_file.seek(0)
        raw_bytes = uploaded_file.read()
        uploaded_file.seek(0)
    except (AttributeError, OSError) as error:
        raise ValueError(
            "The DOCX upload stream is not readable or seekable."
        ) from error

    if not isinstance(raw_bytes, bytes):
        raise ValueError("The DOCX upload did not return binary content.")

    if not raw_bytes:
        raise ValueError("The uploaded DOCX file is empty.")

    return raw_bytes


def _extract_text_blocks(
    document: DocumentObject,
) -> list[DocumentTextBlock]:
    """Extract non-empty paragraphs with heading and section references."""

    blocks: list[DocumentTextBlock] = []
    section_stack: list[str] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = _clean_text(paragraph.text)

        if not text:
            continue

        style = paragraph.style.name if paragraph.style else "Normal"
        heading_level = _heading_level(style)

        if heading_level is not None:
            section_stack = section_stack[: heading_level - 1]
            section_stack.append(text)
            block_type = "heading"
        elif _is_list_style(style):
            block_type = "list_item"
        else:
            block_type = "paragraph"

        blocks.append(
            DocumentTextBlock(
                block_id=f"paragraph-{paragraph_index + 1}",
                paragraph_index=paragraph_index,
                text=text,
                style=style,
                block_type=block_type,
                heading_level=heading_level,
                section_path=list(section_stack),
            )
        )

    return blocks


def _extract_tables(
    document: DocumentObject,
) -> list[DocumentTableData]:
    """Extract all non-empty DOCX tables as normalized DataFrames."""

    extracted: list[DocumentTableData] = []

    for source_index, table in enumerate(document.tables):
        rows = _table_rows(table)

        if not rows:
            continue

        title = _infer_table_title(document, table, source_index)
        dataframe, headers, warnings = _rows_to_dataframe(rows)

        if dataframe.empty and len(dataframe.columns) == 0:
            continue

        extracted_index = len(extracted)
        extracted.append(
            DocumentTableData(
                table_index=extracted_index,
                table_id=f"table-{source_index + 1}",
                title=title,
                row_count=int(len(dataframe)),
                column_count=int(len(dataframe.columns)),
                headers=headers,
                dataframe=dataframe,
                source_reference=f"docx://table/{source_index + 1}",
                warnings=warnings,
            )
        )

    return extracted


def _table_rows(table: Table) -> list[list[str]]:
    """Return cleaned table rows while removing fully empty rows."""

    cleaned_rows: list[list[str]] = []

    for row in table.rows:
        cells = [_clean_text(cell.text) for cell in row.cells]

        if any(cells):
            cleaned_rows.append(cells)

    return cleaned_rows


def _rows_to_dataframe(
    rows: list[list[str]],
) -> tuple[pd.DataFrame, list[str], list[str]]:
    """Convert Word table rows into a DataFrame with unique headers."""

    warnings: list[str] = []
    width = max(len(row) for row in rows)
    normalized_rows = [
        row + [""] * (width - len(row))
        for row in rows
    ]

    first_row = normalized_rows[0]
    use_first_row_as_header = _looks_like_header(
        first_row=first_row,
        remaining_rows=normalized_rows[1:],
    )

    if use_first_row_as_header:
        raw_headers = first_row
        data_rows = normalized_rows[1:]
    else:
        raw_headers = [
            f"column_{index + 1}"
            for index in range(width)
        ]
        data_rows = normalized_rows
        warnings.append(
            "A reliable header row was not detected; generated column "
            "names were used."
        )

    headers, header_warnings = _unique_headers(raw_headers)
    warnings.extend(header_warnings)
    dataframe = pd.DataFrame(data_rows, columns=headers)

    dataframe = dataframe.replace(r"^\s*$", pd.NA, regex=True)
    dataframe = dataframe.dropna(how="all")
    dataframe = dataframe.dropna(axis=1, how="all")
    dataframe = dataframe.reset_index(drop=True)

    for column in dataframe.columns:
        dataframe[column] = _infer_series_type(dataframe[column])

    return dataframe, list(dataframe.columns), warnings


def _looks_like_header(
    first_row: list[str],
    remaining_rows: list[list[str]],
) -> bool:
    """Use conservative rules to decide whether row one is a header."""

    non_empty = [value for value in first_row if value]

    if not non_empty:
        return False

    if len(set(value.casefold() for value in non_empty)) != len(non_empty):
        return False

    alpha_cells = sum(
        bool(re.search(r"[A-Za-z\u1000-\u109F]", value))
        for value in non_empty
    )

    if alpha_cells / len(non_empty) < 0.5:
        return False

    if not remaining_rows:
        return True

    first_numeric = sum(_is_numeric_text(value) for value in first_row)
    later_numeric = sum(
        _is_numeric_text(value)
        for row in remaining_rows[:5]
        for value in row
        if value
    )

    if first_numeric == 0 and later_numeric > 0:
        return True

    common_header_terms = {
        "id",
        "name",
        "date",
        "product",
        "category",
        "region",
        "quantity",
        "price",
        "cost",
        "revenue",
        "profit",
        "metric",
        "value",
        "အမည်",
        "ရက်စွဲ",
        "ကုန်ပစ္စည်း",
        "ဒေသ",
        "အရေအတွက်",
        "တန်ဖိုး",
    }

    tokens = {
        token.casefold()
        for cell in first_row
        for token in re.split(r"\s+|_", cell)
        if token
    }

    return bool(tokens & common_header_terms)


def _unique_headers(
    raw_headers: list[str],
) -> tuple[list[str], list[str]]:
    """Create clean, non-empty, unique DataFrame headers."""

    headers: list[str] = []
    warnings: list[str] = []
    counts: dict[str, int] = {}

    for index, raw_header in enumerate(raw_headers):
        base = _normalize_header(raw_header)

        if not base:
            base = f"column_{index + 1}"
            warnings.append(
                f"Empty header at position {index + 1} was renamed to '{base}'."
            )

        count = counts.get(base, 0) + 1
        counts[base] = count
        header = base if count == 1 else f"{base}_{count}"

        if count > 1:
            warnings.append(
                f"Duplicate header '{base}' was renamed to '{header}'."
            )

        headers.append(header)

    return headers, warnings


def _normalize_header(value: str) -> str:
    """Normalize a Word table header for schema linking."""

    normalized = _clean_text(value).casefold()
    normalized = re.sub(r"[^0-9a-z\u1000-\u109f]+", "_", normalized)
    normalized = re.sub(r"_+", "_", normalized)
    return normalized.strip("_")


def _infer_series_type(series: pd.Series) -> pd.Series:
    """Infer simple numeric or date types without destroying text values."""

    non_null = series.dropna()

    if non_null.empty:
        return series

    numeric = pd.to_numeric(non_null, errors="coerce")

    if numeric.notna().mean() >= 0.95:
        converted = pd.to_numeric(series, errors="coerce")

        if converted.dropna().map(lambda value: float(value).is_integer()).all():
            return converted.astype("Int64")

        return converted.astype("Float64")

    if _looks_like_date_series(non_null):
        parsed = pd.to_datetime(series, errors="coerce")

        if parsed.notna().sum() >= max(1, int(len(non_null) * 0.95)):
            return parsed

    return series.astype("string").str.strip()


def _looks_like_date_series(series: pd.Series) -> bool:
    """Identify common date-like strings before invoking date parsing."""

    samples = series.astype(str).head(20)
    date_pattern = re.compile(
        r"^(?:\d{4}[-/]\d{1,2}[-/]\d{1,2}|"
        r"\d{1,2}[-/]\d{1,2}[-/]\d{2,4})$"
    )
    return samples.map(lambda value: bool(date_pattern.match(value.strip()))).mean() >= 0.8


def _is_numeric_text(value: str) -> bool:
    """Return True for a plain numeric table cell."""

    cleaned = value.replace(",", "").strip()

    if not cleaned:
        return False

    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def _heading_level(style_name: str) -> int | None:
    """Extract heading level from common Word heading style names."""

    match = re.search(r"heading\s*(\d+)", style_name, flags=re.IGNORECASE)

    if match:
        return int(match.group(1))

    if style_name.casefold() == "title":
        return 1

    if style_name.casefold() == "subtitle":
        return 2

    return None


def _is_list_style(style_name: str) -> bool:
    """Return True for common numbered and bulleted list styles."""

    normalized = style_name.casefold()
    return "list" in normalized or "bullet" in normalized


def _infer_table_title(
    document: DocumentObject,
    table: Table,
    table_index: int,
) -> str:
    """Use a nearby preceding paragraph as a human-readable table title."""

    body = document._element.body
    table_position = list(body).index(table._element)

    for element in reversed(list(body)[:table_position]):
        if element.tag.endswith("}p"):
            paragraph = Paragraph(element, document)
            text = _clean_text(paragraph.text)

            if text:
                return text[:120]

    return f"Table {table_index + 1}"


def _extract_metadata(
    document: DocumentObject,
    filename: str,
    text_blocks: list[DocumentTextBlock],
    tables: list[DocumentTableData],
) -> DocxMetadata:
    """Create document-level extraction metadata."""

    properties = document.core_properties
    warnings: list[str] = []

    if not text_blocks:
        warnings.append("The DOCX contains no non-empty narrative paragraphs.")

    if not tables:
        warnings.append("The DOCX contains no usable tables.")

    return DocxMetadata(
        filename=filename,
        format="docx",
        paragraph_count=int(len(document.paragraphs)),
        non_empty_paragraph_count=int(len(text_blocks)),
        heading_count=int(
            sum(block.block_type == "heading" for block in text_blocks)
        ),
        table_count=int(len(document.tables)),
        extracted_table_count=int(len(tables)),
        text_character_count=int(
            sum(len(block.text) for block in text_blocks)
        ),
        title=_clean_optional(properties.title),
        author=_clean_optional(properties.author),
        subject=_clean_optional(properties.subject),
        created=_datetime_text(properties.created),
        modified=_datetime_text(properties.modified),
        warnings=warnings,
    )


def _clean_text(value: str) -> str:
    """Normalize whitespace while preserving English and Myanmar text."""

    text = str(value or "").replace("\xa0", " ")
    text = re.sub(r"[\t\r\n]+", " ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _clean_optional(value: Any) -> str | None:
    """Return a stripped string or None."""

    if value is None:
        return None

    cleaned = str(value).strip()
    return cleaned or None


def _datetime_text(value: Any) -> str | None:
    """Serialize a document date safely."""

    if value is None:
        return None

    try:
        return value.isoformat()
    except AttributeError:
        return str(value)
