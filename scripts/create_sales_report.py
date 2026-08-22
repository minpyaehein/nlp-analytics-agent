"""Create a bilingual sample sales report in DOCX format."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("sample_data/sales_report.docx")
MYANMAR_FONT = "Myanmar Text"
ENGLISH_FONT = "Aptos"


def set_run_font(
    run,
    font_name: str,
    font_size: float = 10,
    bold: bool = False,
) -> None:
    """Apply a font consistently for English and Myanmar text."""

    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.rFonts

    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)

    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), font_name)
    run_fonts.set(qn("w:cs"), font_name)


def set_cell_text(
    cell,
    value,
    *,
    bold: bool = False,
    font_size: float = 8,
) -> None:
    """Put consistently formatted text inside a table cell."""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(str(value))
    set_run_font(
        run,
        font_name=ENGLISH_FONT,
        font_size=font_size,
        bold=bold,
    )

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_formatted_heading(
    document: Document,
    text: str,
    level: int,
    font_size: float,
    centered: bool = False,
):
    """Add and format a document heading."""

    heading = document.add_heading(text, level=level)

    if centered:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in heading.runs:
        set_run_font(
            run,
            font_name=ENGLISH_FONT,
            font_size=font_size,
            bold=True,
        )

    return heading


def create_sales_report() -> Path:
    """Create the bilingual sample sales report and return its path."""

    document = Document()

    document.core_properties.title = "Sample Sales Report"
    document.core_properties.author = "Min Pyae Hein"
    document.core_properties.subject = (
        "Sample bilingual sales dataset for InsightFlow AI"
    )

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    add_formatted_heading(
        document,
        "Sales Performance Report",
        level=1,
        font_size=20,
        centered=True,
    )

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description_run = description.add_run(
        "ရန်ကုန်၊ မန္တလေးနှင့် နေပြည်တော်ဒေသများ၏ "
        "အရောင်းအချက်အလက်များ ဖြစ်ပါသည်။"
    )
    set_run_font(
        description_run,
        font_name=MYANMAR_FONT,
        font_size=11,
    )

    add_formatted_heading(
        document,
        "Sales Transactions",
        level=2,
        font_size=14,
    )

    headers = [
        "order_id",
        "order_date",
        "product",
        "category",
        "region",
        "quantity",
        "unit_price",
        "unit_cost",
    ]

    rows = [
        [1001, "2026-01-05", "Laptop", "Electronics", "Yangon", 2, 850, 700],
        [1002, "2026-01-07", "Mouse", "Accessories", "Mandalay", 5, 20, 12],
        [1003, "2026-02-02", "Keyboard", "Accessories", "Yangon", 3, 45, 30],
        [1004, "2026-02-10", "Monitor", "Electronics", "Naypyidaw", 2, 240, 190],
        [1005, "2026-03-01", "Laptop", "Electronics", "Mandalay", 1, 850, 700],
        [1006, "2026-03-08", "Mouse", "Accessories", "Yangon", 8, 20, 12],
        [1007, "2026-04-03", "Monitor", "Electronics", "Yangon", 3, 240, 190],
        [1008, "2026-04-15", "Keyboard", "Accessories", "Mandalay", 4, 45, 30],
        [1009, "2026-05-06", "Laptop", "Electronics", "Yangon", 2, 850, 700],
        [1010, "2026-05-20", "Mouse", "Accessories", "Naypyidaw", 6, 20, 12],
        [1010, "2026-05-20", "Mouse", "Accessories", "Naypyidaw", 6, 20, 12],
        [1011, "2026-06-03", "Keyboard", "Accessories", "", 2, 45, 30],
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(
            cell,
            header,
            bold=True,
            font_size=7.5,
        )

    for row_data in rows:
        cells = table.add_row().cells

        for cell, value in zip(cells, row_data):
            set_cell_text(
                cell,
                value,
                font_size=8,
            )

    document.add_paragraph()

    add_formatted_heading(
        document,
        "Management Note",
        level=2,
        font_size=14,
    )

    management_note = document.add_paragraph()
    management_run = management_note.add_run(
        "Revenue နှင့် Profit တို့ကို table ရှိ quantity၊ "
        "unit_price နှင့် unit_cost columns များမှ "
        "deterministic calculation ဖြင့် တွက်ချက်ရမည်။"
    )
    set_run_font(
        management_run,
        font_name=MYANMAR_FONT,
        font_size=10.5,
    )

    formula_heading = document.add_paragraph()
    formula_heading_run = formula_heading.add_run(
        "Calculation formulas"
    )
    set_run_font(
        formula_heading_run,
        font_name=ENGLISH_FONT,
        font_size=10.5,
        bold=True,
    )

    revenue_formula = document.add_paragraph(style="List Bullet")
    revenue_run = revenue_formula.add_run(
        "Revenue = quantity × unit_price"
    )
    set_run_font(
        revenue_run,
        font_name=ENGLISH_FONT,
        font_size=10,
    )

    profit_formula = document.add_paragraph(style="List Bullet")
    profit_run = profit_formula.add_run(
        "Profit = quantity × (unit_price - unit_cost)"
    )
    set_run_font(
        profit_run,
        font_name=ENGLISH_FONT,
        font_size=10,
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)

    print(f"Created {OUTPUT_PATH}")
    print(f"Transactions: {len(rows)}")
    print(f"Columns: {len(headers)}")

    return OUTPUT_PATH


if __name__ == "__main__":
    create_sales_report()
